import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import sys
import os
import glob

sys.stdout.reconfigure(encoding='utf-8')

working_dir = os.path.dirname(os.path.abspath(__file__))
logo_file = os.path.join(working_dir, 'atsede_teguhan_logo_hd.png')

print("=" * 70)
print("  አጸደ ትጉሃን ሰንበት ትምህርት ቤት - የሰርተፍኬት ማመንጫ ፕሮግራም")
print("=" * 70)

excel_files = [f for f in glob.glob(os.path.join(working_dir, '*.xlsx')) if not f.endswith('_backup.xlsx') and not 'mail_merge' in f]
if not excel_files:
    print("❌ ምንም የ Excel ፋይል አልተገኘም! እባክዎ የክፍሉን Excel ፋይል እዚህ ፎልደር ውስጥ ያስቀምጡ።")
    sys.exit(1)

src_excel = excel_files[0]
print(f"📄 የተመረጠው የ Excel ፋይል፦ {os.path.basename(src_excel)}")

wb = openpyxl.load_workbook(src_excel, data_only=True)
ws = wb.active

students = []
for r in range(3, ws.max_row + 1):
    name = ws.cell(row=r, column=3).value
    if name is None or not str(name).strip():
        continue
    
    no = ws.cell(row=r, column=2).value
    raw_scores = [ws.cell(row=r, column=c).value for c in range(6, 11)]
    valid_scores = [s for s in raw_scores if isinstance(s, (int, float))]
    if not valid_scores:
        continue
    
    total = sum(valid_scores)
    avg = total / len(valid_scores)
    avg_rounded = round(avg, 2) if avg % 1 != 0 else int(avg)
    
    students.append({
        'row': r,
        'no': str(no).strip() if no else "",
        'name': str(name).strip(),
        'scores': raw_scores,
        's1': raw_scores[0] if raw_scores[0] is not None else "-",
        's2': raw_scores[1] if raw_scores[1] is not None else "-",
        's3': raw_scores[2] if raw_scores[2] is not None else "-",
        's4': raw_scores[3] if raw_scores[3] is not None else "-",
        's5': raw_scores[4] if raw_scores[4] is not None else "-",
        'total': total,
        'avg': avg_rounded,
        'grade': ws.title if ws.title else "3ኛ ክፍል"
    })

if not students:
    print("❌ ምንም የተማሪ ዳታ አልተገኘም!")
    sys.exit(1)

all_totals = [s['total'] for s in students]
for s in students:
    s['rank'] = sum(1 for t in all_totals if t > s['total']) + 1

print(f"✅ {len(students)} ተማሪዎች ተገኝተዋል፤ ደረጃቸውም በትክክል ተሰልቷል!")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    sectPr = section._sectPr
    pgBorders = parse_xml(f'''
        <w:pgBorders {nsdecls("w")} w:offsetFrom="page">
            <w:top w:val="triple" w:sz="18" w:space="12" w:color="8B4513"/>
            <w:left w:val="triple" w:sz="18" w:space="12" w:color="8B4513"/>
            <w:bottom w:val="triple" w:sz="18" w:space="12" w:color="8B4513"/>
            <w:right w:val="triple" w:sz="18" w:space="12" w:color="8B4513"/>
        </w:pgBorders>
    ''')
    sectPr.append(pgBorders)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=90, bottom=90, left=110, right=110):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, color="DAA520", sz="4", val="single"):
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

total_students = len(students)
for idx, s in enumerate(students):
    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(4)
    run0 = p0.add_run("❖  በስመ አብ ወወልድ ወመንፈስ ቅዱስ አሐዱ አምላክ አሜን።  ❖")
    run0.font.name = "Nyala"
    run0.font.size = Pt(12)
    run0.font.bold = True
    run0.font.color.rgb = RGBColor(120, 80, 20)

    hdr_table = doc.add_table(rows=1, cols=3)
    hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    c_left = hdr_table.rows[0].cells[0]
    c_left.width = Inches(1.35)
    p_l = c_left.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(logo_file):
        r_img = p_l.add_run()
        r_img.add_picture(logo_file, width=Inches(1.25))

    c_mid = hdr_table.rows[0].cells[1]
    c_mid.width = Inches(4.8)
    p_m = c_mid.paragraphs[0]
    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_m.paragraph_format.space_after = Pt(2)
    
    r_ch = p_m.add_run("በኢትዮጵያ ኦርቶዶክስ ተዋሕዶ ቤተ ክርስቲያን የአዲስ አበባ ሀገረ ስብከት\n")
    r_ch.font.name = "Nyala"
    r_ch.font.size = Pt(11.5)
    r_ch.font.bold = True
    r_ch.font.color.rgb = RGBColor(70, 70, 70)
    
    r_sc = p_m.add_run("አጸደ ትጉሃን ሰንበት ትምህርት ቤት\n")
    r_sc.font.name = "Nyala"
    r_sc.font.size = Pt(24)
    r_sc.font.bold = True
    r_sc.font.color.rgb = RGBColor(150, 30, 15)
    
    r_ti = p_m.add_run("የተማሪዎች የውጤት መግለጫ እና የምስክር ወረቀት")
    r_ti.font.name = "Nyala"
    r_ti.font.size = Pt(14.5)
    r_ti.font.bold = True
    r_ti.font.underline = True
    r_ti.font.color.rgb = RGBColor(20, 35, 90)

    c_right = hdr_table.rows[0].cells[2]
    c_right.width = Inches(1.3)
    p_r = c_right.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_me = p_r.add_run(f"📅 2018 ዓ.ም\n📚 {s['grade']}")
    r_me.font.name = "Nyala"
    r_me.font.size = Pt(11.5)
    r_me.font.bold = True
    r_me.font.color.rgb = RGBColor(110, 70, 20)

    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep.paragraph_format.space_before = Pt(6)
    p_sep.paragraph_format.space_after = Pt(12)
    r_sep = p_sep.add_run("════════════════════════════════════════════════════════")
    r_sep.font.color.rgb = RGBColor(212, 175, 55)

    info_table = doc.add_table(rows=2, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_cells = [
        ("የተማሪው ሙሉ ስም፦", f"{s['name']}", "ክፍል፦", f"{s['grade']}"),
        ("አጠቃላይ ድምር ውጤት፦", f"{s['total']}", "አማካኝ ነጥብ፦", f"{s['avg']}%")
    ]
    for r_idx, r_data in enumerate(info_cells):
        tbl_row = info_table.rows[r_idx]
        for c_idx, text in enumerate(r_data):
            cell = tbl_row.cells[c_idx]
            set_cell_background(cell, "FDF8F2" if c_idx % 2 == 0 else "FFFFFF")
            set_cell_borders(cell, color="E2C898", sz="4")
            set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
            p = cell.paragraphs[0]
            run = p.add_run(str(text))
            run.font.name = "Nyala"
            run.font.size = Pt(12)
            if c_idx % 2 == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(90, 60, 20)
            else:
                run.font.bold = True
                run.font.color.rgb = RGBColor(10, 30, 110) if c_idx == 1 else RGBColor(0, 0, 0)

    p_rank = doc.add_paragraph()
    p_rank.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rank.paragraph_format.space_before = Pt(14)
    p_rank.paragraph_format.space_after = Pt(14)
    r_rk_l = p_rank.add_run("🏆 ያገኘው/ችው ደረጃ፦  ")
    r_rk_l.font.name = "Nyala"
    r_rk_l.font.size = Pt(14.5)
    r_rk_l.font.bold = True
    r_rk_l.font.color.rgb = RGBColor(120, 60, 10)
    
    r_rk_v = p_rank.add_run(f" [  {s['rank']}ኛ ደረጃ  ] ")
    r_rk_v.font.name = "Nyala"
    r_rk_v.font.size = Pt(18)
    r_rk_v.font.bold = True
    r_rk_v.font.color.rgb = RGBColor(180, 20, 20)

    subjects_data = [
        ("ተ.ቁ", "የትምህርት ዓይነት", "የተገኘ ውጤት", "ከፍተኛ ውጤት"),
        ("1", "መሠረተ እምነት", str(s['s1']), "100"),
        ("2", "ቅዱሳት መጻሕፍት", str(s['s2']), "100"),
        ("3", "ክርስቲያናዊ ሥነ ምግባር", str(s['s3']), "100"),
        ("4", "ሥርዓተ ቤተ ክርስቲያን", str(s['s4']), "100"),
        ("5", "የቤተ ክርስቲያን ታሪክ", str(s['s5']), "100"),
        ("", "አጠቃላይ ድምር እና አማካኝ ውጤት", f"{s['total']}", f"{s['avg']}%")
    ]
    sub_table = doc.add_table(rows=len(subjects_data), cols=4)
    sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for s_idx, s_row in enumerate(subjects_data):
        tbl_row = sub_table.rows[s_idx]
        is_header = (s_idx == 0)
        is_summary = (s_idx == len(subjects_data) - 1)
        for c_idx, val in enumerate(s_row):
            cell = tbl_row.cells[c_idx]
            if is_header:
                set_cell_background(cell, "A04000")
                set_cell_borders(cell, color="78281F", sz="6")
            elif is_summary:
                set_cell_background(cell, "FEF5E7")
                set_cell_borders(cell, color="D4AC0D", sz="6")
            else:
                set_cell_background(cell, "FAFAFA" if s_idx % 2 == 1 else "FFFFFF")
                set_cell_borders(cell, color="E5E7E9", sz="4")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (c_idx != 1) else (WD_ALIGN_PARAGRAPH.RIGHT if is_summary else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(str(val))
            run.font.name = "Nyala"
            run.font.size = Pt(11.5)
            if is_header:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            elif is_summary:
                run.font.bold = True
                run.font.color.rgb = RGBColor(120, 40, 10) if c_idx == 1 else RGBColor(0, 0, 120)
            elif c_idx == 2:
                run.font.bold = True
                run.font.color.rgb = RGBColor(10, 30, 110)

    p_sig_space = doc.add_paragraph()
    p_sig_space.paragraph_format.space_before = Pt(50)
    
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_titles = ["የሰባሳቢው ፊርማ", "የትምህርት ክፍል ፊርማ", "የቤተክርስቲያን አስተዳዳሪ ፊርማ"]
    for c_idx, title in enumerate(sig_titles):
        c0 = sig_table.rows[0].cells[c_idx]
        set_cell_borders(c0, color="FFFFFF", sz="0")
        p_c0 = c0.paragraphs[0]
        p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c0.paragraph_format.space_after = Pt(20)
        run_line = p_c0.add_run("_______________________")
        run_line.font.color.rgb = RGBColor(140, 140, 140)
        
        c1 = sig_table.rows[1].cells[c_idx]
        set_cell_borders(c1, color="FFFFFF", sz="0")
        p_c1 = c1.paragraphs[0]
        p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_c1.add_run(title)
        run_title.font.name = "Nyala"
        run_title.font.size = Pt(11.5)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(60, 60, 60)

    p_footer = doc.add_paragraph()
    p_footer.paragraph_format.space_before = Pt(24)
    r_date = p_footer.add_run("ቀን፦ ጳጉሜ 1 2018 ዓ.ም                                                 የሰንበት ትምህርት ቤቱ ማኅተም")
    r_date.font.name = "Nyala"
    r_date.font.size = Pt(11)
    r_date.font.color.rgb = RGBColor(90, 90, 90)

    if idx < total_students - 1:
        doc.add_page_break()

out_name = os.path.join(working_dir, f"ሰርተፍኬት_የተማሪዎች_ሙሉ.docx")
doc.save(out_name)
print("=" * 70)
print(f"🎉 በስኬት ተጠናቋል! የሰርተፍኬት ፋይል ተፈጥሯል፦ {os.path.basename(out_name)}")
print("=" * 70)
